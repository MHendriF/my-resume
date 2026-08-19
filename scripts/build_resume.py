"""
Professional Resume Generator Pipeline (DOCX, PDF & Markdown)
Supports Multi-Target Role Variants in both English & Indonesian with Clickable Hyperlinks.

Usage:
    python scripts/build_resume.py --target all          # Build all 8 variants (EN + ID)
    python scripts/build_resume.py --target general      # Build Master Resume (EN)
    python scripts/build_resume.py --target general_id   # Build Master Resume (ID)
    python scripts/build_resume.py --target frontend     # Build Frontend Resume (EN)
    python scripts/build_resume.py --target android      # Build Android Resume (EN)
    python scripts/build_resume.py --target web3         # Build Web3 Resume (EN)
"""

import os
import sys
import json
import re
import html
import argparse
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

# Fix Windows console UTF-8 output
if sys.stdout.encoding != 'utf-8':
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except Exception:
        pass

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA_DIR = os.path.join(BASE_DIR, 'data')
TEMPLATES_DIR = os.path.join(DATA_DIR, 'templates')
OUTPUT_DIR = os.path.join(BASE_DIR, 'output')
EN_DIR = os.path.join(OUTPUT_DIR, 'en')
ID_DIR = os.path.join(OUTPUT_DIR, 'id')

os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(EN_DIR, exist_ok=True)
os.makedirs(ID_DIR, exist_ok=True)

# Elegant Color Palette
PRIMARY_COLOR = RGBColor(15, 23, 42)      # Deep Slate #0F172A
SECONDARY_COLOR = RGBColor(37, 99, 235)  # Professional Blue #2563EB
TEXT_COLOR = RGBColor(30, 41, 59)        # Slate Charcoal #1E293B
MUTED_COLOR = RGBColor(71, 85, 105)      # Slate Muted #475569

SECTION_TITLES = {
    "en": {
        "summary": "PROFESSIONAL SUMMARY",
        "experience": "WORK EXPERIENCE",
        "skills": "TECHNICAL SKILLS",
        "education": "EDUCATION",
        "certifications": "CERTIFICATIONS",
        "languages": "LANGUAGES"
    },
    "id": {
        "summary": "RINGKASAN PROFESIONAL",
        "experience": "PENGALAMAN KERJA",
        "skills": "KEAHLIAN TEKNIS",
        "education": "PENDIDIKAN",
        "certifications": "SERTIFIKASI",
        "languages": "KEMAMPUAN BAHASA"
    }
}

def load_json(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        return json.load(f)

def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_hyperlink(paragraph, url, text, hex_color="2563EB", underline=True, font_size_pt=9):
    """Inserts a real, clickable hyperlink into a python-docx paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = parse_xml(f'<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" r:id="{r_id}" w:history="1"/>')
    
    r = parse_xml(r'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    rPr = parse_xml(r'<w:rPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    
    rFonts = parse_xml(r'<w:rFonts xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:ascii="Calibri" w:hAnsi="Calibri"/>')
    rPr.append(rFonts)
    
    sz_val = int(font_size_pt * 2)
    sz = parse_xml(f'<w:sz xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="{sz_val}"/>')
    rPr.append(sz)
    
    c = parse_xml(f'<w:color xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="{hex_color}"/>')
    rPr.append(c)
    
    if underline:
        u = parse_xml(r'<w:u xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="single"/>')
        rPr.append(u)
        
    escaped_text = html.escape(text)
    t = parse_xml(f'<w:t xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xml:space="preserve">{escaped_text}</w:t>')
    
    r.append(rPr)
    r.append(t)
    hyperlink.append(r)
    paragraph._p.append(hyperlink)

def build_docx(profile, skills, template, output_docx_path):
    lang = template.get('lang', 'en')
    titles = SECTION_TITLES.get(lang, SECTION_TITLES['en'])
    
    doc = Document()
    
    # 0.5 Inch Standard ATS Margins
    for section in doc.sections:
        section.top_margin = Inches(0.48)
        section.bottom_margin = Inches(0.48)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    def add_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = SECONDARY_COLOR
        run.font.name = 'Calibri'
        
        # Bottom Accent Border
        pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                         r'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="2563EB"/>'
                         r'</w:pBdr>')
        p._p.get_or_add_pPr().append(pBdr)

    def add_table_header(left_text, right_text, is_sub=False):
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        
        row = table.rows[0]
        c_left, c_right = row.cells[0], row.cells[1]
        c_left.width = Inches(5.3)
        c_right.width = Inches(2.2)
        
        set_cell_margins(c_left, top=0, bottom=0, left=0, right=0)
        set_cell_margins(c_right, top=0, bottom=0, left=0, right=0)
        
        # Left Text
        p_left = c_left.paragraphs[0]
        p_left.paragraph_format.space_before = Pt(2 if not is_sub else 0)
        p_left.paragraph_format.space_after = Pt(1)
        p_left.paragraph_format.line_spacing = 1.05
        
        r_left = p_left.add_run(left_text)
        r_left.font.name = 'Calibri'
        r_left.font.size = Pt(10)
        r_left.bold = not is_sub
        r_left.font.color.rgb = PRIMARY_COLOR if not is_sub else TEXT_COLOR
        
        # Right Text (Date / Location strictly 1 line)
        p_right = c_right.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_right.paragraph_format.space_before = Pt(2 if not is_sub else 0)
        p_right.paragraph_format.space_after = Pt(1)
        p_right.paragraph_format.line_spacing = 1.05
        
        r_right = p_right.add_run(right_text)
        r_right.font.name = 'Calibri'
        r_right.font.size = Pt(9.5)
        r_right.italic = True
        r_right.font.color.rgb = MUTED_COLOR

    def add_bullet(text, is_first=False, is_last=False):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_before = Pt(1 if is_first else 0)
        p.paragraph_format.space_after = Pt(2 if is_last else 0.5)
        p.paragraph_format.line_spacing = 1.08
        
        # Parse **bold** markdown tokens into distinct Word runs
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if not part: continue
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
                run.font.color.rgb = PRIMARY_COLOR
            else:
                run = p.add_run(part)
                run.font.color.rgb = TEXT_COLOR
            run.font.name = 'Calibri'
            run.font.size = Pt(9.5)

    # 1. Header (Name & Clickable Contact Line)
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(1)
    
    r_name = p_name.add_run(profile['name'].upper())
    r_name.bold = True
    r_name.font.size = Pt(18)
    r_name.font.color.rgb = PRIMARY_COLOR
    r_name.font.name = 'Calibri'

    # Clickable Contact Line
    c = profile['contact']
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_before = Pt(0)
    p_contact.paragraph_format.space_after = Pt(4)
    
    # Location
    r_loc = p_contact.add_run(f"{c.get('location', '')}  •  ")
    r_loc.font.size = Pt(9)
    r_loc.font.color.rgb = MUTED_COLOR
    r_loc.font.name = 'Calibri'
    
    # Email (mailto link)
    add_hyperlink(p_contact, f"mailto:{c.get('email', '')}", c.get('email', ''), hex_color="2563EB", underline=False, font_size_pt=9)
    r_sep1 = p_contact.add_run("  •  ")
    r_sep1.font.size = Pt(9)
    r_sep1.font.color.rgb = MUTED_COLOR
    r_sep1.font.name = 'Calibri'

    # Phone (tel link)
    phone_clean = c.get('phone', '').replace(' ', '').replace('-', '')
    add_hyperlink(p_contact, f"tel:{phone_clean}", c.get('phone', ''), hex_color="2563EB", underline=False, font_size_pt=9)
    r_sep2 = p_contact.add_run("  •  ")
    r_sep2.font.size = Pt(9)
    r_sep2.font.color.rgb = MUTED_COLOR
    r_sep2.font.name = 'Calibri'

    # LinkedIn (clickable)
    add_hyperlink(p_contact, c.get('linkedin_url', 'https://linkedin.com/in/mhendrif'), "LinkedIn", hex_color="2563EB", underline=True, font_size_pt=9)
    r_sep3 = p_contact.add_run("  •  ")
    r_sep3.font.size = Pt(9)
    r_sep3.font.color.rgb = MUTED_COLOR
    r_sep3.font.name = 'Calibri'

    # GitHub (clickable)
    add_hyperlink(p_contact, c.get('github_url', 'https://github.com/MHendriF'), "GitHub", hex_color="2563EB", underline=True, font_size_pt=9)
    r_sep4 = p_contact.add_run("  •  ")
    r_sep4.font.size = Pt(9)
    r_sep4.font.color.rgb = MUTED_COLOR
    r_sep4.font.name = 'Calibri'

    # Portfolio & Career Graph (clickable)
    portfolio_label = "Portofolio" if lang == 'id' else "Portfolio"
    add_hyperlink(p_contact, c.get('portfolio_url', 'https://mhendrif.github.io/my-resume/'), portfolio_label, hex_color="2563EB", underline=True, font_size_pt=9)

    # 2. Summary
    summary_key = template.get('summary_key', 'general')
    summary_text = profile.get('summaries', {}).get(summary_key, profile.get('summaries', {}).get('general', ''))
    
    add_heading(titles['summary'])
    p_sum = doc.add_paragraph()
    p_sum.paragraph_format.space_before = Pt(1)
    p_sum.paragraph_format.space_after = Pt(4)
    p_sum.paragraph_format.line_spacing = 1.12
    r_sum = p_sum.add_run(summary_text)
    r_sum.font.size = Pt(9.5)
    r_sum.font.color.rgb = TEXT_COLOR
    r_sum.font.name = 'Calibri'

    # 3. Work Experience
    add_heading(titles['experience'])
    exp_map = {exp['id']: exp for exp in profile.get('experiences', [])}
    ordered_ids = template.get('experience_order', [exp['id'] for exp in profile.get('experiences', [])])
    
    for eid in ordered_ids:
        if eid in exp_map:
            exp = exp_map[eid]
            role_header = f"{exp['role']} | {exp['company']} ({exp['location']})"
            add_table_header(role_header, exp['period'])
            
            bullets = exp.get('bullets', [])
            for idx, bullet in enumerate(bullets):
                add_bullet(bullet, is_first=(idx == 0), is_last=(idx == len(bullets) - 1))

    # 4. Technical Skills
    add_heading(titles['skills'])
    skills_key = template.get('skills_key', 'general')
    target_skills = skills.get(skills_key, skills.get('general', []))
    
    for cat in target_skills:
        p_skill = doc.add_paragraph()
        p_skill.paragraph_format.space_before = Pt(0)
        p_skill.paragraph_format.space_after = Pt(1.5)
        p_skill.paragraph_format.line_spacing = 1.08
        
        r_cat = p_skill.add_run(f"{cat['category']}: ")
        r_cat.bold = True
        r_cat.font.size = Pt(9.5)
        r_cat.font.color.rgb = PRIMARY_COLOR
        r_cat.font.name = 'Calibri'
        
        r_items = p_skill.add_run(cat['items'])
        r_items.font.size = Pt(9.5)
        r_items.font.color.rgb = TEXT_COLOR
        r_items.font.name = 'Calibri'

    # 5. Education
    add_heading(titles['education'])
    for edu in profile.get('education', []):
        add_table_header(f"{edu['institution']} | {edu['degree']}", edu['period'])
        p_edu = doc.add_paragraph()
        p_edu.paragraph_format.space_before = Pt(0)
        p_edu.paragraph_format.space_after = Pt(3)
        gpa_label = "IPK" if lang == 'id' else "GPA"
        focus_label = "Fokus Keahlian" if lang == 'id' else "Focus Area"
        r_edu = p_edu.add_run(f"{gpa_label}: {edu['gpa']}  |  {focus_label}: {edu['focus']}")
        r_edu.font.size = Pt(9)
        r_edu.font.color.rgb = MUTED_COLOR
        r_edu.font.name = 'Calibri'

    # 6. Certifications
    add_heading(titles['certifications'])
    for cert in profile.get('certifications', []):
        add_table_header(f"{cert['title']} | {cert['issuer']}", cert['period'])
        if 'focus' in cert:
            p_cert_det = doc.add_paragraph()
            p_cert_det.paragraph_format.space_before = Pt(0)
            p_cert_det.paragraph_format.space_after = Pt(2)
            p_cert_det.paragraph_format.left_indent = Inches(0.15)
            focus_label = "Kompetensi Kunci" if lang == 'id' else "Core Focus"
            r_det = p_cert_det.add_run(f"{focus_label}: {cert['focus']}")
            r_det.font.size = Pt(9)
            r_det.font.color.rgb = MUTED_COLOR
            r_det.font.name = 'Calibri'

    # 7. Languages
    add_heading(titles['languages'])
    p_lang = doc.add_paragraph()
    p_lang.paragraph_format.space_before = Pt(1)
    p_lang.paragraph_format.space_after = Pt(4)
    for lang_item in profile.get('languages', []):
        r_l = p_lang.add_run(f"{lang_item['language']}: ")
        r_l.bold = True
        r_l.font.size = Pt(9.5)
        r_l.font.color.rgb = PRIMARY_COLOR
        r_l.font.name = 'Calibri'
        
        r_lp = p_lang.add_run(f"{lang_item['proficiency']}    ")
        r_lp.font.size = Pt(9.5)
        r_lp.font.color.rgb = TEXT_COLOR
        r_lp.font.name = 'Calibri'

    os.makedirs(os.path.dirname(output_docx_path), exist_ok=True)
    doc.save(output_docx_path)
    print(f"  [DOCX] -> {output_docx_path}")

def build_markdown(profile, skills, template, output_md_path):
    lang = template.get('lang', 'en')
    titles = SECTION_TITLES.get(lang, SECTION_TITLES['en'])
    c = profile['contact']
    summary_key = template.get('summary_key', 'general')
    summary_text = profile.get('summaries', {}).get(summary_key, profile.get('summaries', {}).get('general', ''))
    skills_key = template.get('skills_key', 'general')
    target_skills = skills.get(skills_key, skills.get('general', []))
    exp_map = {exp['id']: exp for exp in profile.get('experiences', [])}
    ordered_ids = template.get('experience_order', [exp['id'] for exp in profile.get('experiences', [])])

    lines = []
    lines.append(f"# {profile['name']}")
    lines.append(f"{c['location']} | [{c['email']}](mailto:{c['email']}) | [{c['phone']}](tel:{c['phone']})")
    lines.append(f"[LinkedIn: {c['linkedin']}]({c['linkedin_url']}) | [GitHub: {c['github']}]({c['github_url']}) | [Portfolio: {c['portfolio']}]({c['portfolio_url']})\n")
    lines.append("---\n")
    lines.append(f"## {titles['summary']}")
    lines.append(summary_text + "\n")
    lines.append("---\n")
    lines.append(f"## {titles['experience']}\n")

    for eid in ordered_ids:
        if eid in exp_map:
            exp = exp_map[eid]
            lines.append(f"### {exp['role']} | {exp['company']} ({exp['location']})")
            lines.append(f"*{exp['period']}*\n")
            for b in exp.get('bullets', []):
                lines.append(f"- {b}")
            lines.append("")

    lines.append("---\n")
    lines.append(f"## {titles['skills']}\n")
    for cat in target_skills:
        lines.append(f"- **{cat['category']}:** {cat['items']}")
    lines.append("")

    lines.append("---\n")
    lines.append(f"## {titles['education']}\n")
    for edu in profile.get('education', []):
        lines.append(f"### {edu['institution']}")
        lines.append(f"**{edu['degree']}** | *{edu['period']}*")
        gpa_label = "IPK" if lang == 'id' else "GPA"
        focus_label = "Fokus Keahlian" if lang == 'id' else "Focus Area"
        lines.append(f"- **{gpa_label}:** {edu['gpa']}")
        lines.append(f"- **{focus_label}:** {edu['focus']}\n")

    lines.append("---\n")
    lines.append(f"## {titles['certifications']}\n")
    for cert in profile.get('certifications', []):
        lines.append(f"- **{cert['title']}** – {cert['issuer']} *({cert['period']})*")
        if 'focus' in cert:
            focus_label = "Kompetensi Kunci" if lang == 'id' else "Core Focus"
            lines.append(f"  *{focus_label}: {cert['focus']}*")
    lines.append("")

    lines.append("---\n")
    lines.append(f"## {titles['languages']}")
    for lang_item in profile.get('languages', []):
        lines.append(f"- **{lang_item['language']}:** {lang_item['proficiency']}")

    content = "\n".join(lines) + "\n"
    os.makedirs(os.path.dirname(output_md_path), exist_ok=True)
    with open(output_md_path, 'w', encoding='utf-8') as f:
        f.write(content)
    print(f"  [MD]   -> {output_md_path}")

def convert_to_pdf(docx_path, pdf_path):
    try:
        import win32com.client
        word = win32com.client.Dispatch('Word.Application')
        word.Visible = False
        try:
            wdoc = word.Documents.Open(os.path.abspath(docx_path))
            wdoc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
            wdoc.Close()
            print(f"  [PDF]  -> {pdf_path}")
        finally:
            word.Quit()
    except Exception as e:
        print(f"  [PDF Warning] Word COM conversion notice: {e}")

def process_target(target_name, profiles, all_skills):
    template_path = os.path.join(TEMPLATES_DIR, f"{target_name}.json")
    if not os.path.exists(template_path):
        print(f"Error: Template {template_path} not found.")
        return

    template = load_json(template_path)
    lang = template.get('lang', 'en')
    profile = profiles.get(lang, profiles['en'])
    skills = all_skills.get(lang, all_skills['en'])
    
    base_name = template.get('output_filename', f"Resume_{target_name.capitalize()}")
    
    # Destination Paths grouped by language (output/en/ or output/id/)
    target_dir = ID_DIR if lang == 'id' else EN_DIR
    docx_path = os.path.join(target_dir, f"{base_name}.docx")
    pdf_path = os.path.join(target_dir, f"{base_name}.pdf")
    md_path = os.path.join(target_dir, f"{base_name}.md")

    display_title = template.get('display_title', target_name)
    print(f"\n🚀 Generating [{display_title}] ({lang.upper()})...")
    build_docx(profile, skills, template, docx_path)
    build_markdown(profile, skills, template, md_path)
    convert_to_pdf(docx_path, pdf_path)

def main():
    parser = argparse.ArgumentParser(description="Resume Build Pipeline")
    parser.add_argument('--target', default='all', help="Target template: general, general_id, frontend, frontend_id, android, android_id, web3, web3_id, all")
    parser.add_argument('--lang', default='all', choices=['en', 'id', 'all'], help="Language filter")
    args = parser.parse_args()

    # Load English and Indonesian Profile & Skills
    profiles = {
        'en': load_json(os.path.join(DATA_DIR, 'profile.json')),
        'id': load_json(os.path.join(DATA_DIR, 'profile_id.json'))
    }
    all_skills = {
        'en': load_json(os.path.join(DATA_DIR, 'skills.json')),
        'id': load_json(os.path.join(DATA_DIR, 'skills_id.json'))
    }

    if args.target == 'all':
        targets = []
        if args.lang in ['en', 'all']:
            targets.extend(['general', 'frontend', 'android', 'web3'])
        if args.lang in ['id', 'all']:
            targets.extend(['general_id', 'frontend_id', 'android_id', 'web3_id'])
            
        for t in targets:
            process_target(t, profiles, all_skills)
    else:
        process_target(args.target, profiles, all_skills)

    print("\n✅ All target builds completed successfully!")

if __name__ == '__main__':
    main()
